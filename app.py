import streamlit as st
import requests
import pandas as pd
import html
from io import BytesIO
import datetime
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import calendar
import time
import re

# ==========================================
# CONFIGURACIÓN INICIAL Y ESTILOS
# ==========================================
st.set_page_config(page_title="Boletín Mensual - Banxico", layout="wide")

# Inyección de CSS para cambiar el color al Azul Banxico (#00205B)
st.markdown("""
    <style>
    div.stButton > button, div.stDownloadButton > button {
        background-color: #00205B !important;
        color: white !important;
        border: none !important;
    }
    div.stButton > button:hover, div.stDownloadButton > button:hover {
        background-color: #00153D !important;
        color: white !important;
    }
    span[data-baseweb="tag"] {
        background-color: #00205B !important;
        color: white !important;
    }
    </style>
""", unsafe_allow_html=True)

# ==========================================
# FUNCIONES AUXILIARES (BACKEND)
# ==========================================

@st.cache_data(show_spinner="Descargando y procesando datos del BIS...")
def load_data_bis():
    url = "https://www.bis.org/api/document_lists/cbspeeches.json"
    response = requests.get(url)
    response.raise_for_status()
    data = response.json()

    speeches_dict = data.get("list", {})
    rows = []

    for path, speech in speeches_dict.items():
        title = html.unescape(speech.get("short_title", ""))
        date_str = speech.get("publication_start_date", "")
        link = "https://www.bis.org" + path + (".htm" if not path.endswith(".htm") else "")

        # El BPI ya trae "Autor: Título" por defecto
        rows.append({"Date": date_str, "Title": title, "Link": link})

    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"])
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner="Navegando y extrayendo discursos del BBk (Alemania)...")
def load_data_bbk(start_date_str, end_date_str):
    base_url = "https://www.bundesbank.de/action/en/730564/bbksearch"
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
    
    rows = []
    page = 0
    
    while True:
        params = {'sort': 'bbksortdate desc', 'dateFrom': start_date_str, 'dateTo': end_date_str, 'pageNumString': str(page)}
        try:
            response = requests.get(base_url, headers=headers, params=params)
            response.raise_for_status()
        except:
            break 
            
        soup = BeautifulSoup(response.text, 'html.parser')
        items = soup.find_all('li', class_='resultlist__item')
        
        if not items:
            break 
            
        for item in items:
            fecha_tag = item.find('span', class_='metadata__date')
            fecha_str = fecha_tag.text.strip() if fecha_tag else ""
            
            author_tag = item.find('span', class_='metadata__authors')
            author_str = author_tag.text.strip() if author_tag else ""
            if author_str:
                # Separa nombres si vienen juntos (Ej: JoachimNagel -> Joachim Nagel)
                author_str = re.sub(r'([a-z])([A-Z])', r'\1 \2', author_str)
            
            data_div = item.find('div', class_='teasable__data')
            link, titulo = "", ""
            
            if data_div:
                a_tag = data_div.find('a', class_='teasable__link')
                if a_tag:
                    link = a_tag.get('href', '')
                    if link.startswith('/'):
                        link = "https://www.bundesbank.de" + link
                    span_tag = a_tag.find('span', class_='link__label')
                    if span_tag:
                        titulo = span_tag.text.strip()
            
            # Formato Autor: Título
            if author_str and titulo:
                titulo = f"{author_str}: {titulo}"
            
            if fecha_str and titulo:
                rows.append({"Date": fecha_str, "Title": titulo, "Link": link})
                
        if len(items) < 10:
            break
        page += 1
        time.sleep(0.5) 
        
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"], format='%d.%m.%Y', errors='coerce')
        df = df.sort_values("Date", ascending=False)
    return df

@st.cache_data(show_spinner="Navegando y extrayendo discursos del BdE (España)...")
def load_data_bde(start_date_str, end_date_str):
    # Usamos el portal en español para capturar los cargos en idioma nativo
    base_url = "https://www.bde.es/wbe/es/noticias-eventos/actualidad-banco-espana/intervenciones-publicas/"
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
    
    try:
        start_date = datetime.datetime.strptime(start_date_str, '%d.%m.%Y')
    except:
        start_date = datetime.datetime(2000, 1, 1)
        
    rows = []
    page = 1
    
    while True:
        params = {'page': page, 'role': ' ', 'sort': 'DESC', 'limit': 10}
        try:
            response = requests.get(base_url, headers=headers, params=params)
            response.raise_for_status()
        except:
            break 
            
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # El BdE usa fechas en formato DD/MM/YYYY
        date_pattern = re.compile(r'\b(\d{2}/\d{2}/\d{4})\b')
        items = soup.find_all(string=date_pattern)
        
        items_found = 0
        
        for date_node in items:
            fecha_str = date_pattern.search(date_node).group(1)
            
            parent = date_node.parent
            a_tag = None
            
            # Subir niveles para buscar el enlace principal
            for _ in range(6):
                if parent is None:
                    break
                a_tags = parent.find_all('a')
                # Excluir botones de compartir
                a_tag = next((a for a in a_tags if 'compartir' not in a.text.lower() and 'share' not in a.text.lower()), None)
                if a_tag:
                    break
                parent = parent.parent
                
            if a_tag and parent:
                link = a_tag.get('href', '')
                if link.startswith('/'):
                    link = "https://www.bde.es" + link
                
                bloque_texto = parent.get_text(separator=" | ", strip=True)
                partes = [p.strip() for p in bloque_texto.split('|') if p.strip()]
                
                titulo_final = a_tag.text.strip()
                autor = ""
                
                # Deducción heurística del autor basado en posición relativa a la fecha
                try:
                    idx_fecha = partes.index(fecha_str)
                    if len(partes) > idx_fecha + 1:
                        posible_autor = partes[idx_fecha + 1]
                        if posible_autor != titulo_final and len(posible_autor) < 50:
                            autor = posible_autor.replace('.', '').strip()
                except:
                    pass
                
                # Formato Autor: Título
                if autor:
                    titulo_final = f"{autor}: {titulo_final}"
                
                if not any(r['Link'] == link for r in rows):
                    rows.append({"Date": fecha_str, "Title": titulo_final, "Link": link})
                    items_found += 1
                    
        # Control para detener paginación si ya rebasamos la fecha de inicio
        should_break = False
        if rows:
            try:
                last_date = datetime.datetime.strptime(rows[-1]['Date'], '%d/%m/%Y')
                if last_date < start_date:
                    should_break = True
            except:
                pass
                
        if items_found == 0 or should_break:
            break
            
        page += 1
        time.sleep(0.5) 
        
    df = pd.DataFrame(rows)
    if not df.empty:
        df["Date"] = pd.to_datetime(df["Date"], format='%d/%m/%Y', errors='coerce')
        df = df.sort_values("Date", ascending=False)
        
    return df

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '0000EE')
    rPr.append(c)

    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)

    sz = docx.oxml.shared.OxmlElement('w:sz')
    sz.set(docx.oxml.shared.qn('w:val'), '24')
    rPr.append(sz)
    
    szCs = docx.oxml.shared.OxmlElement('w:szCs')
    szCs.set(docx.oxml.shared.qn('w:val'), '24')
    rPr.append(szCs)

    rFonts = docx.oxml.shared.OxmlElement('w:rFonts')
    rFonts.set(docx.oxml.shared.qn('w:ascii'), 'Calibri')
    rFonts.set(docx.oxml.shared.qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)

    t = docx.oxml.shared.OxmlElement('w:t')
    t.text = text
    new_run.append(rPr)
    new_run.append(t)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def generate_word(dataframe, title="Discursos", subtitle=""):
    doc = Document()
    
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_sub.add_run(subtitle)
        run_sub.font.name = 'Calibri'
        run_sub.font.size = Pt(12)

    doc.add_paragraph()

    display_cols = [c for c in dataframe.columns if c != 'Link']
    table = doc.add_table(rows=1, cols=len(display_cols))
    
    hdr_cells = table.rows[0].cells
    for idx, header_text in enumerate(display_cols):
        p = hdr_cells[idx].paragraphs[0]
        run = p.add_run(header_text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.bold = True 

    for index, row in dataframe.iterrows():
        row_cells = table.add_row().cells
        
        date_str = str(row['Date'])[:10]
        
        p_date = row_cells[0].paragraphs[0]
        run_date = p_date.add_run(date_str)
        run_date.font.name = 'Calibri'
        run_date.font.size = Pt(12)
        
        if 'Organismo' in display_cols:
            p_org = row_cells[1].paragraphs[0]
            run_org = p_org.add_run(str(row['Organismo']))
            run_org.font.name = 'Calibri'
            run_org.font.size = Pt(12)
            
            p_title = row_cells[2].paragraphs[0]
            add_hyperlink(p_title, str(row['Title']), str(row['Link']))
        else:
            p_title = row_cells[1].paragraphs[0]
            add_hyperlink(p_title, str(row['Title']), str(row['Link']))

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# ==========================================
# INTERFAZ DE USUARIO (SIDEBAR Y NAVEGACIÓN)
# ==========================================

try:
    st.sidebar.image("logo_banxico.png", use_column_width=True)
except:
    st.sidebar.markdown("### 🏦 BANCO DE MÉXICO")

st.sidebar.markdown("---")
st.sidebar.header("Menú de Navegación")

tipo_doc = st.sidebar.selectbox(
    "Selecciona el Tipo de Documento",
    ["Reportes", "Publicaciones Institucionales", "Investigación", "Discursos"]
)

if tipo_doc == "Discursos":
    organismos = [
        "Todos",
        "BBk (Alemania)", "BdE (España)", "BdF (Francia)", "BM", 
        "BoC (Canadá)", "BoE (Inglaterra)", "BoJ (Japón)", "BPI", 
        "CEF", "ECB (Europa)", "Fed (Estados Unidos)", "FMI", "PBoC (China)"
    ]
elif tipo_doc == "Reportes":
    organismos = ["BID", "BM", "BPI", "CEF", "FEM", "OCDE"]
elif tipo_doc == "Investigación":
    organismos = ["BID", "BM", "BPI", "CEMLA", "FMI", "OCDE"]
elif tipo_doc == "Publicaciones Institucionales":
    organismos = ["BM", "BPI", "CEF", "CEMLA", "FMI", "G20", "OCDE", "OEI"]

organismo_seleccionado = st.sidebar.selectbox("Selecciona el Organismo", organismos)

st.sidebar.markdown("---")
st.sidebar.info("Herramienta de extracción automatizada para la elaboración del boletín mensual.")

# ==========================================
# CONTENIDO PRINCIPAL (MAIN PAGE)
# ==========================================

st.title("Boletín Mensual de Organismos Internacionales")
st.markdown(f"**Explorador de {tipo_doc} - {organismo_seleccionado}**")
st.markdown("---")

# ==========================================
# MÓDULOS DE EXTRACCIÓN
# ==========================================

# MÓDULO: DISCURSOS -> TODOS
if tipo_doc == "Discursos" and organismo_seleccionado == "Todos":
    
    st.subheader("1. Selecciona el Mes y Año")

    anios_str = ["2026", "2025", "2024", "2023", "2022", "2021", "2020"]
    meses_dict = {
        "Enero": 1, "Febrero": 2, "Marzo": 3, "Abril": 4,
        "Mayo": 5, "Junio": 6, "Julio": 7, "Agosto": 8,
        "Septiembre": 9, "Octubre": 10, "Noviembre": 11, "Diciembre": 12
    }

    col1, col2 = st.columns(2)
    with col1:
        meses_seleccionados = st.multiselect("Mes(es)", options=list(meses_dict.keys()), default=[])
    with col2:
        anios_seleccionados = st.multiselect("Año(s)", options=anios_str, default=["2026"])

    buscar = st.button("🔍 Buscar", type="primary")

    if buscar or "todos_df_filtrado" in st.session_state:
        if not meses_seleccionados or not anios_seleccionados:
            st.warning("⚠️ Por favor, selecciona al menos un mes y un año.")
        else:
            meses_num = [meses_dict[m] for m in meses_seleccionados]
            anios_num = [int(a) for a in anios_seleccionados]
            
            # Variables de fechas requeridas para scraping BBk y BdE
            min_month, max_month = min(meses_num), max(meses_num)
            min_year, max_year = min(anios_num), max(anios_num)
            start_date_str = f"01.{min_month:02d}.{min_year}"
            last_day = calendar.monthrange(max_year, max_month)[1]
            end_date_str = f"{last_day:02d}.{max_month:02d}.{max_year}"
            
            dfs_combinados = []
            
            # --- 1. Extraer BPI ---
            with st.spinner("Extrayendo del BPI..."):
                df_bpi = load_data_bis()
                if not df_bpi.empty:
                    mask_bpi = (df_bpi["Date"].dt.year.isin(anios_num)) & (df_bpi["Date"].dt.month.isin(meses_num))
                    df_bpi_fil = df_bpi[mask_bpi].copy()
                    if not df_bpi_fil.empty:
                        df_bpi_fil['Organismo'] = 'BPI'
                        dfs_combinados.append(df_bpi_fil)

            # --- 2. Extraer BBk ---
            with st.spinner("Extrayendo del BBk..."):
                df_bbk = load_data_bbk(start_date_str, end_date_str)
                if not df_bbk.empty:
                    mask_bbk = (df_bbk["Date"].dt.year.isin(anios_num)) & (df_bbk["Date"].dt.month.isin(meses_num))
                    df_bbk_fil = df_bbk[mask_bbk].copy()
                    if not df_bbk_fil.empty:
                        df_bbk_fil['Organismo'] = 'BBk'
                        dfs_combinados.append(df_bbk_fil)

            # --- 3. Extraer BdE ---
            with st.spinner("Extrayendo del BdE..."):
                df_bde = load_data_bde(start_date_str, end_date_str)
                if not df_bde.empty:
                    mask_bde = (df_bde["Date"].dt.year.isin(anios_num)) & (df_bde["Date"].dt.month.isin(meses_num))
                    df_bde_fil = df_bde[mask_bde].copy()
                    if not df_bde_fil.empty:
                        df_bde_fil['Organismo'] = 'BdE'
                        dfs_combinados.append(df_bde_fil)

            # --- Consolidar Todo ---
            if dfs_combinados:
                combined_df = pd.concat(dfs_combinados, ignore_index=True)
                combined_df = combined_df.sort_values("Date", ascending=False)
                combined_df = combined_df[['Date', 'Organismo', 'Title', 'Link']]
            else:
                combined_df = pd.DataFrame()
            
            st.session_state["todos_df_filtrado"] = combined_df

            if len(combined_df) > 0:
                st.subheader("2. Resultados de la búsqueda")
                
                col_mensaje, col_boton = st.columns([3, 1])
                with col_mensaje:
                    str_meses = ", ".join(meses_seleccionados)
                    str_anios = ", ".join(anios_seleccionados)
                    st.success(f"Se encontraron **{len(combined_df)}** discursos en total para **{str_meses} {str_anios}**.")
                
                with col_boton:
                    subtitulo_fechas = f"{str_meses} {str_anios}"
                    word_file = generate_word(combined_df, title="Boletín de Discursos", subtitle=subtitulo_fechas)
                    st.download_button(
                        label="📄 Descargar en Word",
                        data=word_file,
                        file_name=f"discursos_todos_{'_'.join(meses_seleccionados)}_{'_'.join(anios_seleccionados)}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                
                display_df = combined_df.copy()
                display_df["Date"] = display_df["Date"].dt.strftime('%Y-%m-%d')
                display_df["Title"] = display_df.apply(lambda x: f"[{x['Title']}]({x['Link']})", axis=1)
                
                st.markdown(display_df[["Date", "Organismo", "Title"]].to_markdown(index=False), unsafe_allow_html=True)
            else:
                st.warning("No se encontraron discursos para las fechas seleccionadas.")


# MÓDULO: DISCURSOS -> BPI
elif tipo_doc == "Discursos" and organismo_seleccionado == "BPI":
    st.subheader("1. Selecciona el Mes y Año")
    df = load_data_bis()
    anios_disponibles = df["Date"].dt.year.dropna().unique().tolist()
    anios_disponibles.sort(reverse=True)
    anios_str = [str(int(a)) for a in anios_disponibles]

    if "2026" in anios_str:
        anios_str.remove("2026")
        anios_str.insert(0, "2026")

    meses_dict = {"Enero": 1, "Febrero": 2, "Marzo": 3, "Abril": 4, "Mayo": 5, "Junio": 6, "Julio": 7, "Agosto": 8, "Septiembre": 9, "Octubre": 10, "Noviembre": 11, "Diciembre": 12}

    col1, col2 = st.columns(2)
    with col1:
        meses_seleccionados = st.multiselect("Mes(es)", options=list(meses_dict.keys()), default=[])
    with col2:
        anios_seleccionados = st.multiselect("Año(s)", options=anios_str, default=["2026"] if "2026" in anios_str else [])

    buscar = st.button("🔍 Buscar", type="primary")

    if buscar or "bis_df_filtrado" in st.session_state:
        if not meses_seleccionados or not anios_seleccionados:
            st.warning("⚠️ Por favor, selecciona al menos un mes y un año.")
        else:
            meses_num = [meses_dict[m] for m in meses_seleccionados]
            anios_num = [int(a) for a in anios_seleccionados]
            
            mask = (df["Date"].dt.year.isin(anios_num)) & (df["Date"].dt.month.isin(meses_num))
            filtered_df = df[mask]
            st.session_state["bis_df_filtrado"] = filtered_df

            if len(filtered_df) > 0:
                st.subheader("2. Resultados de la búsqueda")
                col_mensaje, col_boton = st.columns([3, 1])
                with col_mensaje:
                    str_meses = ", ".join(meses_seleccionados)
                    str_anios = ", ".join(anios_seleccionados)
                    st.success(f"Se encontraron **{len(filtered_df)}** discursos en **{str_meses} {str_anios}**.")
                with col_boton:
                    subtitulo_fechas = f"{str_meses} {str_anios}"
                    word_file = generate_word(filtered_df, title="BPI Central Bank Speeches", subtitle=subtitulo_fechas)
                    st.download_button(label="📄 Descargar en Word", data=word_file, file_name=f"bpi_speeches_{'_'.join(meses_seleccionados)}_{'_'.join(anios_seleccionados)}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)            
                
                filtered_df_display = filtered_df.copy()
                filtered_df_display["Date"] = filtered_df_display["Date"].dt.strftime('%Y-%m-%d')
                filtered_df_display["Title"] = filtered_df_display.apply(lambda x: f"[{x['Title']}]({x['Link']})", axis=1)
                st.markdown(filtered_df_display[["Date", "Title"]].to_markdown(index=False), unsafe_allow_html=True)
            else:
                st.warning("No hay discursos del BPI para las fechas seleccionadas.")
    else:
        st.info("👆 Selecciona el mes y año arriba y presiona **'Buscar'**.")


# MÓDULO: DISCURSOS -> BBk (Alemania)
elif tipo_doc == "Discursos" and organismo_seleccionado == "BBk (Alemania)":
    st.subheader("1. Selecciona el Mes y Año")
    anios_str = ["2026", "2025", "2024", "2023", "2022", "2021", "2020"]
    meses_dict = {"Enero": 1, "Febrero": 2, "Marzo": 3, "Abril": 4, "Mayo": 5, "Junio": 6, "Julio": 7, "Agosto": 8, "Septiembre": 9, "Octubre": 10, "Noviembre": 11, "Diciembre": 12}

    col1, col2 = st.columns(2)
    with col1:
        meses_seleccionados = st.multiselect("Mes(es)", options=list(meses_dict.keys()), default=[])
    with col2:
        anios_seleccionados = st.multiselect("Año(s)", options=anios_str, default=["2026"])

    buscar = st.button("🔍 Buscar", type="primary")

    if buscar or "bbk_df_filtrado" in st.session_state:
        if not meses_seleccionados or not anios_seleccionados:
            st.warning("⚠️ Por favor, selecciona al menos un mes y un año.")
        else:
            meses_num = [meses_dict[m] for m in meses_seleccionados]
            anios_num = [int(a) for a in anios_seleccionados]
            
            min_month, max_month = min(meses_num), max(meses_num)
            min_year, max_year = min(anios_num), max(anios_num)
            start_date_str = f"01.{min_month:02d}.{min_year}"
            last_day = calendar.monthrange(max_year, max_month)[1]
            end_date_str = f"{last_day:02d}.{max_month:02d}.{max_year}"

            df = load_data_bbk(start_date_str, end_date_str)
            
            if not df.empty:
                mask = (df["Date"].dt.year.isin(anios_num)) & (df["Date"].dt.month.isin(meses_num))
                filtered_df = df[mask]
            else:
                filtered_df = pd.DataFrame()
            
            st.session_state["bbk_df_filtrado"] = filtered_df

            if len(filtered_df) > 0:
                st.subheader("2. Resultados de la búsqueda")
                col_mensaje, col_boton = st.columns([3, 1])
                with col_mensaje:
                    str_meses = ", ".join(meses_seleccionados)
                    str_anios = ", ".join(anios_seleccionados)
                    st.success(f"Se encontraron **{len(filtered_df)}** discursos en **{str_meses} {str_anios}**.")
                with col_boton:
                    subtitulo_fechas = f"{str_meses} {str_anios}"
                    word_file = generate_word(filtered_df, title="BBk Central Bank Speeches", subtitle=subtitulo_fechas)
                    st.download_button(label="📄 Descargar en Word", data=word_file, file_name=f"bbk_speeches_{'_'.join(meses_seleccionados)}_{'_'.join(anios_seleccionados)}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                
                filtered_df_display = filtered_df.copy()
                filtered_df_display["Date"] = filtered_df_display["Date"].dt.strftime('%Y-%m-%d')
                filtered_df_display["Title"] = filtered_df_display.apply(lambda x: f"[{x['Title']}]({x['Link']})", axis=1)
                st.markdown(filtered_df_display[["Date", "Title"]].to_markdown(index=False), unsafe_allow_html=True)
            else:
                st.warning("No hay discursos del BBk para las fechas seleccionadas.")
    else:
        st.info("👆 Selecciona el mes y año arriba y presiona **'Buscar'**.")


# MÓDULO: DISCURSOS -> BdE (España)
elif tipo_doc == "Discursos" and organismo_seleccionado == "BdE (España)":
    
    st.subheader("1. Selecciona el Mes y Año")
    anios_str = ["2026", "2025", "2024", "2023", "2022", "2021", "2020"]
    meses_dict = {"Enero": 1, "Febrero": 2, "Marzo": 3, "Abril": 4, "Mayo": 5, "Junio": 6, "Julio": 7, "Agosto": 8, "Septiembre": 9, "Octubre": 10, "Noviembre": 11, "Diciembre": 12}

    col1, col2 = st.columns(2)
    with col1:
        meses_seleccionados = st.multiselect("Mes(es)", options=list(meses_dict.keys()), default=[])
    with col2:
        anios_seleccionados = st.multiselect("Año(s)", options=anios_str, default=["2026"])

    buscar = st.button("🔍 Buscar", type="primary")

    if buscar or "bde_df_filtrado" in st.session_state:
        if not meses_seleccionados or not anios_seleccionados:
            st.warning("⚠️ Por favor, selecciona al menos un mes y un año.")
        else:
            meses_num = [meses_dict[m] for m in meses_seleccionados]
            anios_num = [int(a) for a in anios_seleccionados]
            
            min_month, max_month = min(meses_num), max(meses_num)
            min_year, max_year = min(anios_num), max(anios_num)
            
            start_date_str = f"01.{min_month:02d}.{min_year}"
            last_day = calendar.monthrange(max_year, max_month)[1]
            end_date_str = f"{last_day:02d}.{max_month:02d}.{max_year}"

            df = load_data_bde(start_date_str, end_date_str)
            
            if not df.empty:
                mask = (df["Date"].dt.year.isin(anios_num)) & (df["Date"].dt.month.isin(meses_num))
                filtered_df = df[mask]
            else:
                filtered_df = pd.DataFrame()
            
            st.session_state["bde_df_filtrado"] = filtered_df

            if len(filtered_df) > 0:
                st.subheader("2. Resultados de la búsqueda")
                col_mensaje, col_boton = st.columns([3, 1])
                with col_mensaje:
                    str_meses = ", ".join(meses_seleccionados)
                    str_anios = ", ".join(anios_seleccionados)
                    st.success(f"Se encontraron **{len(filtered_df)}** discursos en **{str_meses} {str_anios}**.")
                with col_boton:
                    subtitulo_fechas = f"{str_meses} {str_anios}"
                    word_file = generate_word(filtered_df, title="BdE Central Bank Speeches", subtitle=subtitulo_fechas)
                    st.download_button(label="📄 Descargar en Word", data=word_file, file_name=f"bde_speeches_{'_'.join(meses_seleccionados)}_{'_'.join(anios_seleccionados)}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                
                filtered_df_display = filtered_df.copy()
                filtered_df_display["Date"] = filtered_df_display["Date"].dt.strftime('%Y-%m-%d')
                filtered_df_display["Title"] = filtered_df_display.apply(lambda x: f"[{x['Title']}]({x['Link']})", axis=1)
                st.markdown(filtered_df_display[["Date", "Title"]].to_markdown(index=False), unsafe_allow_html=True)
            else:
                st.warning("No hay discursos del BdE para las fechas seleccionadas.")
    else:
        st.info("👆 Selecciona el mes y año arriba y presiona **'Buscar'**.")

# MÓDULOS EN CONSTRUCCIÓN (Placeholder para el resto del menú)
else:
    st.info(f"El extractor de **{tipo_doc}** para **{organismo_seleccionado}** está en construcción.")
    st.write("Próximamente podrás extraer estos documentos de forma automatizada.")
